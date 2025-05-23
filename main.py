import telebot
from telebot import types
from my_tokens import tg_token as t, coze_token as c
from cozepy import Coze, TokenAuth, Message, ChatStatus, MessageObjectString
from docx import Document
from os import remove
from docx2pdf import convert
from speech_recognition import Recognizer, WavFile
from audio_extract import extract_audio
from pypdf import PdfReader
import json

bot = telebot.TeleBot(t)
coze = Coze(auth=TokenAuth(c))
states = dict()
r = Recognizer()

# !!!  Сделай папки docx_data, jpg_data, mp3_data, oga_data, pdf_data, png_fata, wav_data
# в каждой папке будут храниться файлы нужного типа данных

@bot.message_handler(content_types=['text', 'document', 'audio', 'voice', 'photo'])
def starting_messages(message):
    print(message)
    # добавление нового пользователя на момент данной сессии
    if message.from_user.id not in states.keys():
        states[message.from_user.id] = [0, "", "", False] # шаг, тип выходных, расширение? , флаг для 2 шага

    # стартовое сообщение
    if states[message.from_user.id][0] == 0 or message.text == "/start":
        if message.text == "/start" and states[message.from_user.id][0] == 0:
            bot.send_message(message.from_user.id, "Привет!\nЯ помогаю сократить содержимое файлов и текстовых сообщений")
            states[message.from_user.id][0] = 1
        elif states[message.from_user.id][0] < 1:
            bot.send_message(message.from_user.id, "Я тебя не понимаю. Напиши /start")

    # добавление файла
    if states[message.from_user.id][0] == 1:
        txt = ''
        f_id = ''
        if message.content_type == 'document':
            f_id = message.document.file_id
        elif message.content_type == 'audio':
            f_id = message.audio.file_id
        elif message.content_type == 'voice':
            f_id = message.voice.file_id
        elif message.content_type == 'photo':
            f_id = message.photo[-1].file_id
        elif message.content_type == 'text' and message.text != "/start":
            txt = message.text

        if txt != '' or f_id != '':
            f = True
            if f_id != '':
                f_obj = bot.get_file(f_id)
                ext = f_obj.file_path[f_obj.file_path.find('.') + 1:]
                if ext in {'docx', 'pdf', 'oga', 'mp3', 'wav', 'jpg', 'png'}:
                    states[message.from_user.id][2] = "f_pt" + f_obj.file_path
                else:
                    f = False
                    bot.send_message(message.from_user.id, "Добавь входные данные в виде текстового или голосового сообщения, или в формате docx, pdf, wav, mp3")
            else:
                states[message.from_user.id][2] = txt
            if f:
                print(txt, f_id)
                states[message.from_user.id][0] = 2
        else:
            bot.send_message(message.from_user.id, "Добавь входные данные в виде текстового или голосового сообщения, или в формате docx, pdf, wav, mp3")

    if states[message.from_user.id][0] == 2:
        button_doc = types.KeyboardButton(text='docx')
        button_pdf = types.KeyboardButton(text='pdf')
        button_txt = types.KeyboardButton(text='Текстовое сообщение')
        keyboard = types.ReplyKeyboardMarkup(one_time_keyboard=True)
        button = [button_txt, button_doc, button_pdf]
        keyboard.add(*button)
        if message.text in ['docx', 'pdf', 'Текстовое сообщение'] and states[message.from_user.id][3]:
            bot.send_message(message.from_user.id, 'Уже работаю над вашим текстом...')
            states[message.from_user.id][0] = 3
            states[message.from_user.id][3] = False
        else:
            states[message.from_user.id][3] = True
            bot.send_message(message.from_user.id, "В каком формате вывести выходные данные?", reply_markup=keyboard)

    # отправка и получение данных от козы
    if states[message.from_user.id][0] == 3:

        telebot.types.ReplyKeyboardRemove()
        states[message.from_user.id][1] = message.text
        print(states[message.from_user.id][0], states[message.from_user.id][1], message.from_user.id)
        if (states[message.from_user.id][2][:4] == 'f_pt'):
            s = states[message.from_user.id][2]
            ext = s[s.rfind('.') + 1:]
            print(ext)
            file = bot.download_file(states[message.from_user.id][2][4:])
            f_name = f'{ext}_data/' + str(message.from_user.id) + '.' + ext
            with open(f_name, 'wb') as new:
                new.write(file)
            new.close()
            if(ext in {'jpg', 'png'}):
                f = coze.files.upload(file)
                chat_poll = coze.chat.create_and_poll(bot_id='7434590500508418104', user_id='0', additional_messages=[Message.build_user_question_objects([MessageObjectString.build_file(f.id)])])
            elif(ext == 'pdf'):
                reader = PdfReader(f_name)
                kol_pages = len(reader.pages)
                states[message.from_user.id][2] = ''
                for i in range(0, kol_pages):
                    page = reader.pages[i]
                    text = page.extract_text()
                    states[message.from_user.id][2] += '\n' + text
                #текст в pdf
            elif (ext == 'docx'):
                doc = Document(f_name)
                fullText = []
                for para in doc.paragraphs:
                    fullText.append(para.text)
                states[message.from_user.id][2] = '\n'.join(fullText)
                print(states[message.from_user.id][2])
                # текст в docx
            elif ext in {'oga', 'mp3', 'wav'}:
                f_name_new = str(message.from_user.id) + '.wav'
                print(f_name_new)
                if ext != 'wav':
                    extract_audio(input_path=f_name, output_path=f_name_new, output_format='wav')
                sample = WavFile(f_name_new)
                with sample as audio:
                    audio = r.record(audio)
                states[message.from_user.id][2] = r.recognize_google(audio, language='ru-Ru')
                remove(f_name_new)
                print(states[message.from_user.id][2])
                #Speech recognition
            else:
                chat_poll = []

        if (states[message.from_user.id][2][:4] != 'f_pt'):
            print('send text to ai')
            chat_poll = coze.chat.create_and_poll(bot_id='7434590500508418104', user_id='0', additional_messages=[Message.build_user_question_text('Сократи текст: ' + states[message.from_user.id][2])])

        answer = []
        for message2 in chat_poll.messages:
            answer.append(message2.content)
        print(answer)
        # если ответ пришел отправка ответа в нужном формате и
        if chat_poll.chat.status == ChatStatus.COMPLETED:
            remove_files = []
            if states[message.from_user.id][2][:4] != 'f_pt':
                data = str(answer[0])
            else:
                a = json.loads(answer[1])['data']['results'][0]['words']
                print(a)
                data = ' '.join([el['text'] for el in a])

            if states[message.from_user.id][1] == "Текстовое сообщение":
                bot.send_message(message.from_user.id, data)

            if states[message.from_user.id][1] == "pdf":
                document = Document()
                document.add_paragraph(data)
                document.add_page_break()
                document_name = f'docx_data/docx{message.from_user.id}.docx'
                document_name1 = f'pdf_data/pdf{message.from_user.id}.pdf'
                document.save(document_name)
                convert(document_name, document_name1)
                file = open(document_name1, 'rb')
                bot.send_document(message.from_user.id, file)
                file.close()
                remove_files = [document_name, document_name1]

            if states[message.from_user.id][1] == "docx":
                document = Document()
                document.add_paragraph(data)
                document.add_page_break()
                document_name = f'docx_data/docx{message.from_user.id}.docx'
                document.save(document_name)
                file = open(document_name, 'rb')
                bot.send_document(message.from_user.id, file)
                file.close()
                remove_files = [document_name]

            #получение ответа от нейросети
            bot.send_message(message.from_user.id, 'Если еще надо чем-то помочь загрузи входные данные')
            for remove_file in remove_files:
                remove(remove_file)
            states[message.from_user.id][0] = 1
        else:
            bot.send_message(message.from_user.id, 'Что-то пошло не так, мы уже решаем эту проблему')


bot.polling(none_stop=True, interval=0)
